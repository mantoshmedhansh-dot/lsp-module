#!/usr/bin/env python3
"""Convert markdown to Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

def create_document():
    doc = Document()

    # Set document properties
    core_props = doc.core_properties
    core_props.title = "CJDQuick OMS/WMS - Complete Process Flow Documentation"
    core_props.author = "CJDQuick System"

    # Title
    title = doc.add_heading('CJDQuick OMS/WMS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete Process Flow Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Meta info
    meta = doc.add_paragraph()
    meta.add_run('Version: ').bold = True
    meta.add_run('1.0\n')
    meta.add_run('Date: ').bold = True
    meta.add_run('January 30, 2026\n')
    meta.add_run('System: ').bold = True
    meta.add_run('Order Management System (OMS) & Warehouse Management System (WMS)')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Executive Summary',
        '2. System Architecture Overview',
        '3. Inbound Process Flow',
        '4. Order-to-Delivery Process Flow',
        '5. Detailed Stage Documentation',
        '6. Material Flow Diagram',
        '7. Status Transitions',
        '8. API Endpoints Reference',
        '9. Integration Points',
        'Appendix A: Glossary',
        'Appendix B: Sample Order Journey'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph(
        'The CJDQuick OMS/WMS system provides end-to-end order fulfillment capabilities '
        'from inventory receipt to proof of delivery. The system follows a structured workflow '
        'ensuring inventory accuracy, order traceability, and operational efficiency.'
    )

    doc.add_heading('Key Principles', 2)
    principles = [
        ('GRN is the First Source of Truth', 'No inventory exists without Goods Receipt'),
        ('FIFO/FEFO Allocation', 'Inventory allocation follows valuation methods'),
        ('Real-time Tracking', 'Every status change is logged and traceable'),
        ('Multi-channel Support', 'Orders from Website, Amazon, Flipkart, Manual entry')
    ]
    for title, desc in principles:
        p = doc.add_paragraph()
        p.add_run(f'{title} - ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # 2. System Architecture
    doc.add_heading('2. System Architecture Overview', 1)

    doc.add_heading('2.1 Module Structure', 2)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'OMS'
    hdr[1].text = 'WMS'
    hdr[2].text = 'B2B Logistics'
    hdr[3].text = 'B2C Courier'
    row = table.rows[1].cells
    row[0].text = 'This Module'
    row[1].text = 'This Module'
    row[2].text = 'Separate'
    row[3].text = 'Separate'

    doc.add_heading('2.2 Technology Stack', 2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Component', 'Technology', 'Purpose']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    tech_data = [
        ('Frontend', 'Next.js 16', 'User Interface'),
        ('Backend', 'FastAPI (Python)', 'API & Business Logic'),
        ('Database', 'PostgreSQL (Supabase)', 'Data Storage'),
        ('Authentication', 'NextAuth + JWT', 'Security')
    ]
    for i, (comp, tech, purpose) in enumerate(tech_data, 1):
        table.rows[i].cells[0].text = comp
        table.rows[i].cells[1].text = tech
        table.rows[i].cells[2].text = purpose

    doc.add_page_break()

    # 3. Inbound Process Flow
    doc.add_heading('3. Inbound Process Flow', 1)

    doc.add_heading('3.1 Inbound Sources', 2)
    doc.add_paragraph(
        'All inventory enters the system through one of these channels:'
    )
    sources = [
        'External PO (Purchase Order) - From vendors/suppliers',
        'ASN (Advance Shipping Notice) - Pre-announced shipments',
        'Stock Transfer (STO) - Inter-warehouse transfers',
        'Sales Returns - Customer returns'
    ]
    for s in sources:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('3.2 GRN Workflow States', 2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Status', 'Description', 'Actions Available']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    grn_states = [
        ('DRAFT', 'GRN created, items being added', 'Edit, Delete, Start Receiving'),
        ('RECEIVING', 'Physical verification in progress', 'Update quantities, Post'),
        ('POSTED', 'Inventory created in system', 'View, Print, Reverse (Admin)'),
        ('CANCELLED', 'GRN cancelled before posting', 'View only')
    ]
    for i, (status, desc, actions) in enumerate(grn_states, 1):
        table.rows[i].cells[0].text = status
        table.rows[i].cells[1].text = desc
        table.rows[i].cells[2].text = actions

    doc.add_page_break()

    # 4. Order-to-Delivery Process Flow
    doc.add_heading('4. Order-to-Delivery Process Flow', 1)

    doc.add_heading('4.1 Complete Order Lifecycle', 2)
    stages = [
        'ORDER CREATED - Customer places order',
        'CONFIRMED - Order validated and verified',
        'ALLOCATED - Inventory reserved for order',
        'PICKLIST GENERATED - Pick instructions created',
        'PICKING - Items being collected from bins',
        'PICKED - All items collected',
        'PACKING - Items being packed',
        'PACKED - Package ready for shipping',
        'INVOICED - Tax invoice generated',
        'MANIFESTED - Added to carrier manifest',
        'DISPATCHED - Handed to carrier',
        'IN TRANSIT - Moving to destination',
        'OUT FOR DELIVERY - With delivery agent',
        'DELIVERED - POD captured, order complete'
    ]
    for i, stage in enumerate(stages, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. ').bold = True
        p.add_run(stage)

    doc.add_heading('4.2 Process Summary', 2)
    table = doc.add_table(rows=12, cols=4)
    table.style = 'Table Grid'
    headers = ['Stage', 'System', 'Output', 'Duration']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    process_data = [
        ('Order Creation', 'OMS', 'Order Record', 'Immediate'),
        ('Confirmation', 'OMS', 'Confirmed Order', '0-24 hrs'),
        ('Allocation', 'WMS', 'Reserved Inventory', 'Immediate'),
        ('Picklist', 'WMS', 'Pick Instructions', 'Immediate'),
        ('Picking', 'WMS', 'Picked Items', '5-30 mins'),
        ('Packing', 'WMS', 'Packed Shipment', '5-15 mins'),
        ('Invoice', 'OMS', 'Tax Invoice', 'Immediate'),
        ('Manifest', 'OMS', 'Carrier Handover', 'Per Batch'),
        ('Dispatch', 'Logistics', 'AWB/Tracking', 'Immediate'),
        ('Last Mile', 'Courier', 'Delivery Attempt', '1-7 days'),
        ('POD', 'Courier', 'Proof of Delivery', 'At Delivery')
    ]
    for i, row_data in enumerate(process_data, 1):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val

    doc.add_page_break()

    # 5. Detailed Stage Documentation
    doc.add_heading('5. Detailed Stage Documentation', 1)

    stages_detail = [
        ('5.1 Stage 1: Order Creation',
         'Capture customer order information from various channels.',
         ['Order Number auto-generated', 'Customer details captured', 'Items with prices added', 'Payment method recorded'],
         'CREATED'),
        ('5.2 Stage 2: Order Confirmation',
         'Validate order and prepare for fulfillment.',
         ['Address validation', 'Payment verification', 'Fraud check', 'SKU availability check'],
         'CONFIRMED'),
        ('5.3 Stage 3: Inventory Allocation',
         'Reserve inventory for the order using FIFO/FEFO logic.',
         ['Find available inventory', 'Apply valuation method', 'Reserve from specific bins', 'Update inventory quantities'],
         'ALLOCATED'),
        ('5.4 Stage 4: Picklist Generation',
         'Create picking instructions for warehouse staff.',
         ['Generate picklist number', 'List bin locations', 'Show required quantities', 'Assign to picker (optional)'],
         'PICKLIST_GENERATED'),
        ('5.5 Stage 5: Picking Process',
         'Physically collect items from warehouse bins.',
         ['Navigate to bin', 'Scan bin barcode', 'Pick required quantity', 'Confirm picked items'],
         'PICKED'),
        ('5.6 Stage 6: Packing Process',
         'Pack items securely for shipping.',
         ['Verify picked items', 'Select packaging', 'Weigh package', 'Generate shipping label'],
         'PACKED'),
        ('5.7 Stage 7: Invoice Generation',
         'Create tax-compliant invoice (GST).',
         ['Invoice number generated', 'GST calculated (CGST/SGST/IGST)', 'HSN codes applied', 'E-way bill (if needed)'],
         'INVOICED'),
        ('5.8 Stage 8: Manifest Creation',
         'Group shipments for carrier handover.',
         ['Manifest number generated', 'AWB numbers listed', 'COD amounts totaled', 'Route assigned'],
         'MANIFESTED'),
        ('5.9 Stage 9: Dispatch',
         'Handover shipments to carrier.',
         ['Carrier pickup', 'Package scan', 'Driver signature', 'Gate pass generated'],
         'DISPATCHED'),
        ('5.10 Stage 10: In Transit',
         'Track shipment movement to delivery.',
         ['Hub scans', 'Transit updates', 'Estimated delivery', 'Exception handling'],
         'IN_TRANSIT'),
        ('5.11 Stage 11: Proof of Delivery',
         'Confirm successful delivery.',
         ['Digital signature', 'OTP verification', 'Photo POD', 'GPS coordinates'],
         'DELIVERED')
    ]

    for title, desc, steps, status in stages_detail:
        doc.add_heading(title, 2)
        p = doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run(desc)

        doc.add_paragraph('Key Steps:', style='List Bullet')
        for step in steps:
            doc.add_paragraph(step, style='List Bullet')

        p = doc.add_paragraph()
        p.add_run('Status: ').bold = True
        p.add_run(status)

    doc.add_page_break()

    # 6. Material Flow
    doc.add_heading('6. Material Flow Diagram', 1)
    doc.add_paragraph(
        'The physical flow of materials through the warehouse follows this path:'
    )

    flow_steps = [
        'INBOUND DOCK - Goods received from supplier',
        'QC AREA - Quality inspection (if required)',
        'PUTAWAY - Move to storage location',
        'STORAGE ZONES - Bulk, Rack, Pick Face, Cold',
        'PICK PATH - Items collected for orders',
        'PACKING STATION - Items packed',
        'STAGING AREA - Ready for dispatch',
        'OUTBOUND DOCK - Handover to carrier'
    ]
    for i, step in enumerate(flow_steps, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. ').bold = True
        p.add_run(step)

    doc.add_heading('Bin Location Structure', 2)
    doc.add_paragraph('Location Code Format: ZONE-AISLE-RACK-LEVEL')
    doc.add_paragraph('Example: A-01-02-03 = Zone A, Aisle 1, Rack 2, Level 3')

    doc.add_page_break()

    # 7. Status Transitions
    doc.add_heading('7. Status Transitions', 1)

    table = doc.add_table(rows=16, cols=3)
    table.style = 'Table Grid'
    headers = ['From Status', 'To Status', 'Trigger']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    transitions = [
        ('CREATED', 'CONFIRMED', 'confirm()'),
        ('CREATED', 'CANCELLED', 'cancel()'),
        ('CREATED', 'ON_HOLD', 'hold()'),
        ('CONFIRMED', 'ALLOCATED', 'allocate()'),
        ('ALLOCATED', 'PICKLIST_GENERATED', 'generatePicklist()'),
        ('PICKLIST_GENERATED', 'PICKING', 'startPicking()'),
        ('PICKING', 'PICKED', 'completePicking()'),
        ('PICKED', 'PACKED', 'pack()'),
        ('PACKED', 'INVOICED', 'generateInvoice()'),
        ('INVOICED', 'DISPATCHED', 'dispatch()'),
        ('DISPATCHED', 'IN_TRANSIT', 'carrierUpdate()'),
        ('IN_TRANSIT', 'OUT_FOR_DELIVERY', 'carrierUpdate()'),
        ('OUT_FOR_DELIVERY', 'DELIVERED', 'confirmDelivery()'),
        ('OUT_FOR_DELIVERY', 'DELIVERY_FAILED', 'deliveryFailed()'),
        ('DELIVERY_FAILED', 'RTO_INITIATED', 'initiateRTO()')
    ]
    for i, (from_s, to_s, trigger) in enumerate(transitions, 1):
        table.rows[i].cells[0].text = from_s
        table.rows[i].cells[1].text = to_s
        table.rows[i].cells[2].text = trigger

    doc.add_page_break()

    # 8. API Endpoints
    doc.add_heading('8. API Endpoints Reference', 1)

    doc.add_heading('8.1 Order Management', 2)
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    headers = ['Method', 'Endpoint', 'Description']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    order_apis = [
        ('POST', '/api/v1/orders', 'Create new order'),
        ('GET', '/api/v1/orders', 'List orders'),
        ('GET', '/api/v1/orders/{id}', 'Get order details'),
        ('PATCH', '/api/v1/orders/{id}', 'Update order'),
        ('POST', '/api/v1/orders/{id}/confirm', 'Confirm order'),
        ('POST', '/api/v1/orders/{id}/allocate', 'Allocate inventory'),
        ('POST', '/api/v1/orders/{id}/cancel', 'Cancel order'),
        ('POST', '/api/v1/orders/{id}/hold', 'Hold order')
    ]
    for i, (method, endpoint, desc) in enumerate(order_apis, 1):
        table.rows[i].cells[0].text = method
        table.rows[i].cells[1].text = endpoint
        table.rows[i].cells[2].text = desc

    doc.add_heading('8.2 Warehouse Operations', 2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Method', 'Endpoint', 'Description']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    wms_apis = [
        ('POST', '/api/v1/picklists', 'Generate picklists'),
        ('GET', '/api/v1/waves/picklists', 'List picklists'),
        ('POST', '/api/v1/waves/{id}/start', 'Start picking'),
        ('POST', '/api/v1/waves/{id}/complete', 'Complete picking'),
        ('POST', '/api/v1/packing/pack', 'Pack order')
    ]
    for i, (method, endpoint, desc) in enumerate(wms_apis, 1):
        table.rows[i].cells[0].text = method
        table.rows[i].cells[1].text = endpoint
        table.rows[i].cells[2].text = desc

    doc.add_page_break()

    # 9. Integration Points
    doc.add_heading('9. Integration Points', 1)

    doc.add_heading('9.1 External Integrations', 2)
    integrations = [
        ('Marketplaces', 'Amazon, Flipkart, Myntra - Order sync'),
        ('Logistics', 'Delhivery, BlueDart, Ecom Express - Shipping'),
        ('Payment', 'Razorpay, PayU - Payment processing'),
        ('ERP', 'SAP, Tally - Accounting sync'),
        ('SMS/Email', 'MSG91, SendGrid - Notifications')
    ]
    for name, desc in integrations:
        p = doc.add_paragraph()
        p.add_run(f'{name}: ').bold = True
        p.add_run(desc)

    doc.add_heading('9.2 Webhook Events', 2)
    events = [
        'order.created - New order received',
        'order.confirmed - Order confirmed',
        'order.allocated - Inventory reserved',
        'order.shipped - Order dispatched',
        'order.delivered - POD received',
        'inventory.received - GRN posted',
        'inventory.adjusted - Stock adjusted'
    ]
    for event in events:
        doc.add_paragraph(event, style='List Bullet')

    doc.add_page_break()

    # Appendix A: Glossary
    doc.add_heading('Appendix A: Glossary', 1)

    table = doc.add_table(rows=13, cols=2)
    table.style = 'Table Grid'
    headers = ['Term', 'Definition']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    glossary = [
        ('AWB', 'Air Waybill - Shipping tracking number'),
        ('COD', 'Cash on Delivery'),
        ('FEFO', 'First Expired, First Out'),
        ('FIFO', 'First In, First Out'),
        ('GRN', 'Goods Receipt Note'),
        ('HSN', 'Harmonized System Nomenclature (Tax code)'),
        ('LIFO', 'Last In, First Out'),
        ('OMS', 'Order Management System'),
        ('POD', 'Proof of Delivery'),
        ('RTO', 'Return to Origin'),
        ('SKU', 'Stock Keeping Unit'),
        ('WMS', 'Warehouse Management System')
    ]
    for i, (term, defn) in enumerate(glossary, 1):
        table.rows[i].cells[0].text = term
        table.rows[i].cells[1].text = defn

    doc.add_page_break()

    # Appendix B: Sample Order Journey
    doc.add_heading('Appendix B: Sample Order Journey', 1)

    doc.add_paragraph('Order: ORD-20260130-NKIT')
    doc.add_paragraph('Customer: Raam')
    doc.add_paragraph('Product: AQP-001 (Qty: 1)')
    doc.add_paragraph('Amount: Rs. 14,160')

    table = doc.add_table(rows=12, cols=4)
    table.style = 'Table Grid'
    headers = ['Timestamp', 'Status', 'Action', 'User']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    journey = [
        ('2026-01-30 11:08', 'CREATED', 'Order placed', 'System'),
        ('2026-01-30 11:51', 'ALLOCATED', 'Inventory reserved', 'Admin'),
        ('2026-01-30 11:54', 'PICKLIST_GENERATED', 'PL-000001 created', 'Admin'),
        ('2026-01-30 12:00', 'PICKING', 'Picker assigned', 'Warehouse'),
        ('2026-01-30 12:15', 'PICKED', 'Items collected', 'Warehouse'),
        ('2026-01-30 12:30', 'PACKED', 'Package ready', 'Packer'),
        ('2026-01-30 12:35', 'INVOICED', 'INV-001 generated', 'System'),
        ('2026-01-30 14:00', 'DISPATCHED', 'Handed to courier', 'System'),
        ('2026-01-30 18:00', 'IN_TRANSIT', 'At carrier hub', 'Carrier'),
        ('2026-01-31 09:00', 'OUT_FOR_DELIVERY', 'With delivery agent', 'Carrier'),
        ('2026-01-31 11:30', 'DELIVERED', 'POD captured', 'Carrier')
    ]
    for i, row_data in enumerate(journey, 1):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('Generated by CJDQuick OMS/WMS System')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.italic = True

    return doc


if __name__ == '__main__':
    doc = create_document()
    output_path = '/Users/mantosh/CJDQuickApp/oms/docs/OMS_WMS_Process_Flow.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
